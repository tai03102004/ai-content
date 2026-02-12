import streamlit as st
import logging
import os
import warnings
from dotenv import load_dotenv
import concurrent.futures
import pdfplumber

# Load environment variables
load_dotenv()

# Import RAG Core mới
from rag_core import get_rag_core

# Suppress warnings
warnings.filterwarnings('ignore', category=UserWarning, message='.*torch.classes.*')

from langchain_text_splitters import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from docling.document_converter import DocumentConverter, PdfFormatOption
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import PdfPipelineOptions, TableFormerMode
from docling.backend.pypdfium2_backend import PyPdfiumDocumentBackend
from docx import Document as DocxDocument

import multiprocessing
from concurrent.futures import ProcessPoolExecutor

# Set protobuf environment variable for efficiency
os.environ["PROTOCOL_BUFFERS_PYTHON_IMPLEMENTATION"] = "python"

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(message)s")
logger = logging.getLogger(__name__)

def process_question(question: str, collection_name: str):
    """
    Xử lý câu hỏi sử dụng RAG Core mới
    """
    rag_core = get_rag_core()
    
    # Gọi hàm answer_question mới
    result = rag_core.answer_question(
        question=question,
        collection_name=collection_name
    )

    logger.info(f"📤 Result: {result}")
    
    if result['success']:
        # Format câu trả lời kèm nguồn (nếu có)
        answer = result['answer']
        sources = result.get('sources', [])
        
        if sources:
            # Loại bỏ trùng lặp và format đẹp
            unique_sources = list(set(sources))
            source_str = "\n".join([f"- *{s}*" for s in unique_sources])
            return f"{answer}\n\n**Nguồn tham khảo:**\n{source_str}"
        return answer
    else:
        return f"⚠️ Lỗi: {result.get('error', 'Unknown error')}"

def process_pdf_file(file_path: str):    
    """
    Hàm xử lý file (PDF/DOCX) -> Chunks & Images
    (Giữ nguyên logic cũ vì nó xử lý tốt việc parse file)
    """
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        chunks = []
        pdf_pages = []

        if file_ext == '.docx':
            docx = DocxDocument(file_path)
            full_text = "\n".join([p.text for p in docx.paragraphs])
            if not full_text.strip(): return [], []
            
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
            doc = Document(
                page_content=full_text,
                metadata={"source": os.path.basename(file_path), "file_path": file_path}
            )
            chunks = text_splitter.split_documents([doc])
            
        else:
            # 1. Configure Docling for PDF
            pipeline_options = PdfPipelineOptions(do_table_structure=True)
            pipeline_options.table_structure_options.do_cell_matching = True
            pipeline_options.table_structure_options.mode = TableFormerMode.ACCURATE

            converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: PdfFormatOption(
                        pipeline_options=pipeline_options,
                        backend=PyPdfiumDocumentBackend
                    )
                }
            )
            
            # 2. Convert PDF
            result = converter.convert(file_path)
            full_markdown = result.document.export_to_markdown()

            if not full_markdown.strip(): return [], []

            # Lấy ảnh trang đầu để preview
            try:
                with pdfplumber.open(file_path) as pdf:
                    if len(pdf.pages) > 0:
                        pdf_pages = [pdf.pages[0].to_image().original]
            except Exception:
                pass

            # 3. Splitting with Markdown awareness
            headers_to_split_on = [("#", "Header 1"), ("##", "Header 2"), ("###", "Header 3")]
            markdown_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers_to_split_on)
            md_header_splits = markdown_splitter.split_text(full_markdown)
            
            # Split tiếp nếu chunk quá lớn
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
            chunks = text_splitter.split_documents(md_header_splits)
            

        # Update Metadata chuẩn cho từng chunk
        for chunk in chunks:
            chunk.metadata.update({
                "source": os.path.basename(file_path),
                "file_path": file_path
            })
        
        return chunks, pdf_pages
        
    except Exception as e:
        logger.error(f"Error processing file {file_path}: {e}")
        # Trả về rỗng để không crash luồng chính
        return [], []

def get_existing_sources(collection_name: str) -> set:
    """Lấy danh sách file đã upload vào collection"""
    existing = set()
    
    try:
        rag_core = get_rag_core()

        import re
        safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', collection_name.lower())

        # Scroll qua tất cả points để lấy unique sources
        offset = None
        while True:
            results, offset = rag_core.qdrant_client.scroll(
                collection_name=safe_name,
                limit=100,
                offset=offset,
                with_payload=True,
                with_vectors=False
            )
            
            for point in results:
                source = point.payload.get("metadata", {}).get("source", "")
                if source:
                    existing.add(source)
            
            if offset is None:
                break
                
    except Exception as e:
        logger.warning(f"⚠️ Could not get existing sources: {e}")
    
    return existing

def ingest_folder(folder_path: str, collection_name: str):
    """
    Hàm upload folder mới: Tương thích với add_documents_smart
    """
    rag_core = get_rag_core()
    
    # 1. Quét file
    files_to_process = []
    if os.path.isdir(folder_path):
        for root, _, files in os.walk(folder_path):
            for file in files:
                if not file.startswith('.') and file.lower().endswith(('.pdf', '.docx')):
                    full_path = os.path.join(root, file)
                    files_to_process.append(full_path)
    
    if not files_to_process:
        st.warning("⚠️ Không tìm thấy file PDF/DOCX nào!")
        return
    
    existing_sources = get_existing_sources(collection_name)

    new_files = []
    skipped_files = []
    
    for f in files_to_process:
        fname = os.path.basename(f)
        if fname in existing_sources:
            skipped_files.append(fname)
        else:
            new_files.append(f)
    
    # Hiển thị files đã skip
    if skipped_files:
        st.info(f"⏭️ Bỏ qua {len(skipped_files)} file đã upload: {', '.join(skipped_files[:5])}...")
        for f in skipped_files:
            if f not in st.session_state["processed_files"]:
                st.session_state["processed_files"].append(f)
    
    if not new_files:
        st.success("✅ Tất cả files đã được upload trước đó!")
        st.session_state["has_db"] = True
        return
    
    st.info(f"📥 Đang xử lý {len(new_files)} file mới...")

    # 2. Xử lý file (Parsing) song song
    total_chunks = []
    all_pages = []
    processed_count = 0
    
    status_bar = st.progress(0, text="🚀 Đang đọc tài liệu...")

    cpu_count = min(multiprocessing.cpu_count(), 8)
    
    with ProcessPoolExecutor(max_workers=cpu_count) as executor:
        future_to_file = {executor.submit(process_pdf_file, f): f for f in new_files} 
        
        for i, future in enumerate(concurrent.futures.as_completed(future_to_file)):
            file_path = future_to_file[future]
            try:
                chunks, pages = future.result()
                if chunks:
                    total_chunks.extend(chunks)
                    if len(all_pages) < 10: 
                        all_pages.extend(pages)
                    
                    file_name = os.path.basename(file_path)
                    if file_name not in st.session_state["processed_files"]:
                        st.session_state["processed_files"].append(file_name)
                    
                    processed_count += 1
            except Exception as e:
                logger.error(f"Lỗi file {file_path}: {e}")
            
            progress = (i + 1) / len(new_files)
            status_bar.progress(progress, text=f"Đang đọc {i+1}/{len(new_files)} files")

    # 3. Upload vào Qdrant (SỬ DỤNG LOGIC MỚI)
    if total_chunks:
        with st.spinner(f"💾 Đang tạo Vector & Upload ({len(total_chunks)} chunks)..."):
            try:
                # [QUAN TRỌNG] Tách content và metadata riêng để gửi vào hàm mới
                texts = [doc.page_content for doc in total_chunks]
                metadatas = [doc.metadata for doc in total_chunks]
                
                # Gọi hàm thông minh của RAGCore (tự batching, tự hybrid, tự retry)
                rag_core.add_documents_smart(
                    collection_name=collection_name, 
                    texts=texts, 
                    metadatas=metadatas
                )
                
                # Cập nhật session preview ảnh
                if "pdf_pages" not in st.session_state: 
                    st.session_state["pdf_pages"] = []
                st.session_state["pdf_pages"].extend(all_pages)
                
                st.success(f"✅ Hoàn tất! Đã xử lý {processed_count} files, tạo {len(total_chunks)} chunks.")
            except Exception as e:
                st.error(f"❌ Lỗi Upload: {e}")
    else:
        st.warning("⚠️ Không trích xuất được nội dung nào từ các file.")
    
    status_bar.empty()

def main() -> None:
    st.set_page_config(page_title="Enterprise RAG v2", page_icon="🏢", layout="wide")
    
    st.subheader("🏢 Enterprise RAG Playground (Hybrid Search)", divider="rainbow")
    col1, col2 = st.columns([1.5, 2])

    # Init Session State
    if "org_name" not in st.session_state: st.session_state["org_name"] = "default"
    if "messages" not in st.session_state: st.session_state["messages"] = []
    if "processed_files" not in st.session_state: st.session_state["processed_files"] = []
    if "pdf_pages" not in st.session_state: st.session_state["pdf_pages"] = []
    if "has_db" not in st.session_state: st.session_state["has_db"] = False

    with col1:
        st.markdown("### 🏢 Organization Setup")
        org_name = st.text_input(
            "Organization Name",
            value=st.session_state["org_name"],
            placeholder="example_corp",
            help="Tên tổ chức sẽ được dùng làm tên Collection trong Qdrant"
        )

        if org_name != st.session_state["org_name"]:
            st.session_state["org_name"] = org_name
            st.session_state["messages"] = []
            st.session_state["processed_files"] = []
            st.session_state["pdf_pages"] = []
            st.session_state["has_db"] = False
            st.rerun()

        COLLECTION_NAME = f"org_{org_name.lower().replace(' ', '_')}"
        import re
        COLLECTION_NAME = re.sub(r'[^a-zA-Z0-9_-]', '_', COLLECTION_NAME)
        if len(COLLECTION_NAME) < 3:
            COLLECTION_NAME = "org_default"
        
        st.info(f"🔗 Collection ID: `{COLLECTION_NAME}`")

        rag_core = get_rag_core()
        try:
            info = rag_core.qdrant_client.get_collection(COLLECTION_NAME)
            st.metric("📊 Documents", f"{info.points_count} chunks")
            if info.points_count > 0:
                st.session_state["has_db"] = True
        except:
            st.metric("📊 Documents", "0 chunks")

        st.markdown("### 📁 Data Ingestion")
        folder_path = st.text_input(
            "Đường dẫn thư mục tài liệu (PDF/DOCX)",
            placeholder="/path/to/documents",
            key="folder_input"
        )
        
        if st.button("🚀 Bắt đầu Xử lý & Upload", type="primary"):
            if folder_path and os.path.exists(folder_path):
                ingest_folder(folder_path, COLLECTION_NAME)
                st.session_state["has_db"] = True
            else:
                st.error("❌ Thư mục không tồn tại!")

        # List Files
        if st.session_state["processed_files"]:
            with st.expander(f"📚 Files đã index ({len(st.session_state['processed_files'])})"):
                for f in st.session_state["processed_files"]:
                    st.text(f"• {f}")

        # Preview Images
        if st.session_state["pdf_pages"]:
            st.markdown("---")
            st.caption("📄 Document Preview")
            zoom = st.slider("Zoom Preview", 100, 800, 300)
            with st.container(height=300, border=True):
                for img in st.session_state["pdf_pages"]:
                    st.image(img, width=zoom)

        st.markdown("---")
        st.markdown("### ⚙️ Database Tools")
        
        if st.button("🗑️ Xóa Dữ liệu (Reset Collection)"):
            rag_core = get_rag_core()
            try:
                # Xóa collection cũ
                rag_core.qdrant_client.delete_collection(COLLECTION_NAME)
                
                # Reset UI state
                st.session_state["processed_files"] = []
                st.session_state["pdf_pages"] = []
                st.session_state["messages"] = []
                st.session_state["has_db"] = False
                
                st.success(f"✅ Đã xóa sạch dữ liệu của '{COLLECTION_NAME}'")
                st.rerun()
            except Exception as e:
                st.error(f"❌ Lỗi: {e}")

    # Chat interface
    with col2:
        st.markdown("### 💬 Hỏi đáp Thông minh")
        msg_container = st.container(height=650, border=True)

        for msg in st.session_state["messages"]:
            avatar = "🤖" if msg["role"] == "assistant" else "👤"
            with msg_container.chat_message(msg["role"], avatar=avatar):
                st.markdown(msg["content"])

        if prompt := st.chat_input("Nhập câu hỏi của bạn..."):
            st.session_state["messages"].append({"role": "user", "content": prompt})
            with msg_container.chat_message("user", avatar="👤"):
                st.markdown(prompt)

            # Logic trả lời
            with msg_container.chat_message("assistant", avatar="🤖"):
                with st.spinner("🔍 Đang tìm kiếm (Hybrid Search) & Suy luận..."):
                    # Luôn cho phép hỏi, nếu chưa có DB thì RAGCore sẽ tự check và trả về rỗng/lỗi nhẹ
                    response = process_question(prompt, COLLECTION_NAME)
                    st.markdown(response)
                
            st.session_state["messages"].append({"role": "assistant", "content": response})

if __name__ == "__main__":
    main()